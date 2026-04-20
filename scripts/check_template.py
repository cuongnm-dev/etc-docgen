#!/usr/bin/env python3
"""Inspect HDSD template: margins, run-level formatting, styles."""

from docx import Document

doc = Document("src/etc_docgen/assets/templates/huong-dan-su-dung.docx")
for i, sec in enumerate(doc.sections):
    top = sec.top_margin
    bot = sec.bottom_margin
    left = sec.left_margin
    right = sec.right_margin
    pw = sec.page_width
    ph = sec.page_height
    print(f"Section {i}:")
    print(f"  Page: {pw / 36000:.0f}x{ph / 36000:.0f} mm")
    print(
        f"  Margins: top={top / 36000:.0f}mm bot={bot / 36000:.0f}mm left={left / 36000:.0f}mm right={right / 36000:.0f}mm"
    )

# Check run-level formatting of key template paragraphs
print()
print("=== Run-level formatting in key paragraphs ===")
for idx in [39, 41, 44, 46, 50, 55, 59, 61, 63, 68, 70]:
    p = doc.paragraphs[idx]
    txt = p.text.strip()[:60]
    print(f"[{idx}] {txt!r}")
    for r in p.runs:
        sz = r.font.size
        sz_pt = f"{sz.pt:.0f}pt" if sz else "None"
        print(
            f'       b={r.bold!s:5s} i={r.italic!s:5s} sz={sz_pt:6s} fn={r.font.name!s:20s} "{r.text[:40]}"'
        )

# Check which paragraphs have direct alignment override
print()
print("=== Paragraphs with direct alignment override ===")
for i, p in enumerate(doc.paragraphs):
    if p.alignment is not None:
        print(f"[{i}] align={p.alignment!s} style={p.style.name!r} text={p.text[:50]!r}")
