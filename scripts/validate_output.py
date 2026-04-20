#!/usr/bin/env python3
"""Validate HDSD output formatting after style fix."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document("examples/minimal/out/huong-dan-su-dung.docx")

# 1. Check margins
print("=== Margins ===")
sec = doc.sections[0]
print(
    f"  top={sec.top_margin / 36000:.0f}mm  bot={sec.bottom_margin / 36000:.0f}mm  left={sec.left_margin / 36000:.0f}mm  right={sec.right_margin / 36000:.0f}mm"
)

# 2. Check ETC_Content style
print("\n=== ETC_Content style ===")
style = doc.styles["ETC_Content"]
print(f"  align={style.paragraph_format.alignment!s}")
print(f"  font={style.font.name} size={style.font.size.pt if style.font.size else 'None'}pt")

# 3. Check heading styles
print("\n=== Heading styles ===")
for name in ["A_HEADING 1", "A_Heading 2", "A_Heading 3", "A_Heading 4"]:
    s = doc.styles[name]
    sz = f"{s.font.size.pt:.0f}pt" if s.font.size else "None"
    print(
        f"  {name:20s} size={sz:5s} bold={s.font.bold!s:5s} italic={s.font.italic!s:5s} align={s.paragraph_format.alignment!s}"
    )

# 4. Check actual paragraph alignment in output
print("\n=== Sample paragraph alignment (rendered) ===")
checks = [
    (21, "purpose"),
    (23, "scope"),
    (37, "description"),
    (41, "preconditions"),
    (50, "step"),
]
for idx, label in checks:
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        # effective alignment: paragraph-level override or style
        direct_align = p.alignment
        style_align = p.style.paragraph_format.alignment if p.style else None
        effective = direct_align if direct_align is not None else style_align
        print(
            f"  [{idx}] {label:15s} direct={direct_align!s:8s} style={style_align!s:10s} effective={effective!s}"
        )

# 5. Count alignment distribution
print("\n=== Alignment distribution in content paragraphs ===")
counts = {"JUSTIFY": 0, "CENTER": 0, "LEFT": 0, "RIGHT": 0, "None/inherit": 0}
for p in doc.paragraphs:
    if p.style and p.style.name == "ETC_Content":
        a = p.alignment
        if a is None:
            counts["None/inherit"] += 1
        elif a == WD_ALIGN_PARAGRAPH.JUSTIFY:
            counts["JUSTIFY"] += 1
        elif a == WD_ALIGN_PARAGRAPH.CENTER:
            counts["CENTER"] += 1
        else:
            counts["LEFT"] += 1
for k, v in counts.items():
    if v > 0:
        print(f"  ETC_Content {k}: {v}")

# 6. First feature structure
print("\n=== First feature [36-61] ===")
for i in range(36, 62):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        txt = p.text.strip()
        if txt:
            sz = None
            for r in p.runs:
                if r.font.size:
                    sz = r.font.size.pt
                    break
            style_sz = p.style.font.size.pt if p.style and p.style.font.size else None
            print(f"  [{i}] style={p.style.name!r:20s} sz={sz or style_sz!s:5s} text={txt[:80]!r}")

print(f"\nTotal: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
