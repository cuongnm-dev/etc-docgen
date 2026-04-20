#!/usr/bin/env python3
"""Fix HDSD template footer: replace legacy text with page number right-aligned."""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TEMPLATE = "src/etc_docgen/assets/templates/huong-dan-su-dung.docx"

doc = Document(TEMPLATE)
sec = doc.sections[0]
footer = sec.footer

# Clear all existing paragraphs in footer
for p in footer.paragraphs:
    p._element.getparent().remove(p._element)

# Create new paragraph with right-aligned PAGE field
p_elem = OxmlElement("w:p")

# Paragraph properties: right-align
pPr = OxmlElement("w:pPr")
jc = OxmlElement("w:jc")
jc.set(qn("w:val"), "right")
pPr.append(jc)
p_elem.append(pPr)


# Run properties (Times New Roman 11pt)
def make_rPr():
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")  # 11pt = 22 half-points
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "22")
    rPr.append(szCs)
    return rPr


# PAGE field: fldChar begin → instrText → fldChar separate → fldChar end
r1 = OxmlElement("w:r")
r1.append(make_rPr())
fc_begin = OxmlElement("w:fldChar")
fc_begin.set(qn("w:fldCharType"), "begin")
r1.append(fc_begin)
p_elem.append(r1)

r2 = OxmlElement("w:r")
r2.append(make_rPr())
instr = OxmlElement("w:instrText")
instr.set(qn("xml:space"), "preserve")
instr.text = " PAGE "
r2.append(instr)
p_elem.append(r2)

r3 = OxmlElement("w:r")
r3.append(make_rPr())
fc_sep = OxmlElement("w:fldChar")
fc_sep.set(qn("w:fldCharType"), "separate")
r3.append(fc_sep)
p_elem.append(r3)

r4 = OxmlElement("w:r")
r4.append(make_rPr())
t = OxmlElement("w:t")
t.text = "1"
r4.append(t)
p_elem.append(r4)

r5 = OxmlElement("w:r")
r5.append(make_rPr())
fc_end = OxmlElement("w:fldChar")
fc_end.set(qn("w:fldCharType"), "end")
r5.append(fc_end)
p_elem.append(r5)

# Append to footer
footer._element.append(p_elem)

doc.save(TEMPLATE)
print("Done: footer = page number, right-aligned, Times New Roman 11pt")

# Verify
doc2 = Document(TEMPLATE)
f2 = doc2.sections[0].footer
p2 = f2.paragraphs[0]
has_page_field = any(ic.text and "PAGE" in ic.text for ic in p2._element.iter(qn("w:instrText")))
print(f"Verify: text={p2.text!r}, has PAGE field={has_page_field}")
